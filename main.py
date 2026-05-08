import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.io as pio
from openai import OpenAI
import os
import re
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import io
import base64
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn 

# ==========================================
# 1. 页面配置与 UI 样式 (保持不变)
# ==========================================
st.set_page_config(page_title="飘哥的智能分析平台", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    .main-title {
        color: #1e3a8a; font-family: "Microsoft YaHei", sans-serif;
        font-weight: 800; border-bottom: 3px solid #1e3a8a;
        padding-bottom: 5px; margin-bottom: 10px;
    }
    [data-testid="stForm"] {
        background-color: #ffffff !important;
        padding: 15px !important;
        border-radius: 8px !important;
        border: 1px solid #e2e8f0 !important;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05) !important;
    }
    button[data-baseweb="tab"] p { font-size: 20px !important; font-weight: 600 !important; }
    .stWidgetLabel p { font-size: 18px !important; font-weight: 500 !important; color: #000000 !important; margin-bottom: 8px !important; }
    span[data-baseweb="tag"] { background-color: #f1f5f9 !important; border: 1px solid #cbd5e1 !important; color: #1e293b !important; }
    .report-box {
        background-color: #ffffff; padding: 25px; border-left: 8px solid #1e3a8a;
        border-radius: 4px; line-height: 1.8; color: #1f2937;
        font-family: "SimSun", "STSong", serif; font-size: 15px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05); margin-top: 15px;
    }
    .block-container { padding-top: 1.5rem !important; }
    hr { margin-top: 0.5rem !important; margin-bottom: 0.5rem !important; }
    div.stButton > button:first-child[kind="primary"] {
        background-color: #1e3a8a !important; color: #ffffff !important; border: none !important;
    }
    .stExpander { border: 1px solid #e2e8f0 !important; background-color: #ffffff !important; }
    .table-title { font-size: 20px; font-weight: 700; color: #1e3a8a; margin: 15px 0 10px 0; display: flex; align-items: center; }
    .plotly-graph-div { background-color: white !important; }
    </style>
    """, unsafe_allow_html=True)

if "last_analysis" not in st.session_state: st.session_state.last_analysis = ""
if "report_list" not in st.session_state: st.session_state.report_list = [] 
if "report_title" not in st.session_state: st.session_state.report_title = "数据智能研判报告"
if "custom_prompt" not in st.session_state:
    st.session_state.custom_prompt = """你是一位深耕上海住建领域、精通网格化管理与12345热线数据治理的首席专家。
你的任务是根据提供上海新建商品房热线投诉的统计数据进行深度研判。内容需包含“现状特征”、“原因分析”、“治理建议”，引用数值，400-600字。"""

# ==========================================
# 2. 核心工具函数 (保持不变)
# ==========================================
def get_system_font():
    font_paths = ["C:/Windows/Fonts/msyh.ttc", "/System/Library/Fonts/PingFang.ttc",
                  "/System/Library/Fonts/STHeiti Light.ttc", "C:/Windows/Fonts/simhei.ttf"]
    for path in font_paths:
        if os.path.exists(path): return path
    return None

def generate_professional_wordcloud(df, label_col, value_col):
    word_freq = dict(zip(df[label_col].astype(str), df[value_col]))
    f_path = get_system_font()
    wc = WordCloud(font_path=f_path, width=1200, height=600, background_color='white', colormap='Blues').generate_from_frequencies(word_freq)
    fig, ax = plt.subplots(figsize=(15, 7))
    ax.imshow(wc, interpolation='bilinear'); ax.axis("off")
    return fig

def generate_html_report(title, items):
    now = pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')
    html_head = f"""<html><head><meta charset="utf-8"><title>{title}</title>
    <script src="https://cdn.plot.ly/plotly-2.24.1.min.js"></script>
    <style>
        body {{ font-family: 'Microsoft YaHei', sans-serif; background: #f4f7f9; padding: 40px; }}
        .container {{ max-width: 1000px; margin: auto; background: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ color: #1e3a8a; border-bottom: 3px solid #1e3a8a; padding-bottom: 10px; }}
        .section {{ margin-bottom: 40px; padding: 20px; border-bottom: 1px solid #eee; }}
        .analysis-text {{ background: #f9fafb; padding: 20px; border-left: 8px solid #1e3a8a; line-height: 1.8; color: #333; }}
        .plotly-graph-div {{ background-color: white !important; }}
    </style></head><body><div class="container"><h1>{title}</h1><p>生成时间：{now}</p>"""
    body = ""
    for item in items:
        safe_text = item['text'].replace('\n', '<br>')
        body += f"<div class='section'><h2>{item['title']}</h2>"
        if item.get("fig"):
            if hasattr(item["fig"], "to_json"):
                item["fig"].update_layout(template="plotly_white", paper_bgcolor='white')
                body += pio.to_html(item["fig"], full_html=False, include_plotlyjs=False)
            else:
                buf = io.BytesIO()
                item["fig"].savefig(buf, format="png", bbox_inches='tight')
                img_str = base64.b64encode(buf.getvalue()).decode()
                body += f"<img src='data:image/png;base64,{img_str}' style='width:100%;'>"
        body += f"<div class='analysis-text'>{safe_text}</div></div>"
    return html_head + body + "</div></body></html>"

def generate_word_report(title, items):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name, style.font.size = 'Times New Roman', Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(10)
    h = doc.add_heading(title, 0)
    for run in h.runs:
        run.font.name, run.font.color.rgb = 'Times New Roman', None
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    doc.add_paragraph(f"生成时间：{pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
    for item in items:
        h1 = doc.add_heading(item['title'], level=1)
        for run in h1.runs:
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if item.get("fig"):
            img_stream = io.BytesIO()
            try:
                if hasattr(item["fig"], "to_image"):
                    img_bytes = item["fig"].to_image(format="png", width=1000, height=550)
                    img_stream.write(img_bytes)
                else:
                    item["fig"].savefig(img_stream, format="png", bbox_inches='tight')
                img_stream.seek(0)
                doc.add_picture(img_stream, width=Inches(6))
            except Exception as e: doc.add_paragraph(f"[图片导出失败: {e}]")
        clean_text = item['text'].replace('**', '') 
        for para_text in clean_text.split('\n'):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.first_line_indent = Inches(0.3)
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.add_page_break()
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def get_sh_expert_analysis(df_shown, group_col, method, client, model_id, sys_prompt, env_type):
    data_summary = df_shown.to_string(index=False)
    user_prompt = f"【分析任务】上海住建运行指标深度研判\n【维度】: {group_col} | 【方法】: {method}\n【数据样本】:\n{data_summary}"
    payload = {"model": model_id, "messages": [{"role": "system", "content": sys_prompt}, {"role": "user", "content": user_prompt}], "temperature": 0.4}
    if not env_type: payload["extra_body"] = {"chat_template_kwargs": {"enable_thinking": False}}
    try:
        response = client.chat.completions.create(**payload)
        content = response.choices[0].message.content
        content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL)
        return content.strip()
    except Exception as e: return f"研判引擎暂时无法连接: {e}"

# ==========================================
# 3. 侧边栏控制 (保持不变)
# ==========================================
with st.sidebar:
    st.markdown("## 🏗️ 系统控制台")
    env_type = st.toggle("🌐 切换至阿里云模式", value=False)
    if env_type:
        default_url, default_key = "https://dashscope.aliyuncs.com/compatible-mode/v1", "sk-819cb281e7c44980a4115f7698b46a1f"
        preset_models = ["qwen-turbo","qwen-long","qwen-plus","qwen-max", "自定义"]
    else:
        default_url, default_key = "http://172.16.25.247:7861/v1", "weiqu"
        preset_models = ["Qwen3.5-35B", "自定义"]

    local_url = st.text_input("🔗 接口地址", value=default_url)
    local_key = st.text_input("🔑 API Key", value=default_key, type="password")
    selected_model = st.selectbox("🤖 选择/输入模型", options=preset_models)
    model_id = st.text_input("✍️ 模型 ID", value=selected_model) if selected_model == "自定义" else selected_model
    client = OpenAI(api_key=local_key, base_url=local_url)

# ==========================================
# 4. 主业务流程
# ==========================================
st.markdown('<h1 class="main-title">🏢 飘哥的数据智能分析平台</h1>', unsafe_allow_html=True)
uploaded_file = st.file_uploader("📤 载入数据文件", type=["xlsx", "xls", "csv"], label_visibility="collapsed")

if uploaded_file:
    if "current_filename" not in st.session_state or st.session_state.current_filename != uploaded_file.name:
        st.session_state.selected_cols, st.session_state.last_analysis, st.session_state.report_list = [], "", []
        st.session_state.current_filename = uploaded_file.name
        st.rerun()

    try:
        df_raw = pd.read_csv(uploaded_file) if uploaded_file.name.endswith('.csv') else pd.read_excel(uploaded_file)
        all_cols = df_raw.columns.tolist()
        if not st.session_state.get("selected_cols"): st.session_state.selected_cols = all_cols

        tab1, tab2 = st.tabs(["🔍 数据筛选与预览", "📊 报表中心与专家研判"])

        with tab1:
            with st.expander("🛠️ 字段管理", expanded=True):
                with st.form("field_config_form"):
                    safe_defaults = [c for c in st.session_state.selected_cols if c in all_cols]
                    temp_selected = st.multiselect("选取分析字段：", all_cols, default=safe_defaults)
                    btn_col1, btn_col2, _ = st.columns([1.2, 1.2, 5])
                    if btn_col1.form_submit_button("🗑️ 一键清空"): st.session_state.selected_cols = []; st.rerun()
                    if btn_col2.form_submit_button("🚀 应用字段", type="primary"): st.session_state.selected_cols = temp_selected; st.rerun()

            st.markdown("---")
            display_cols, df_final = st.session_state.selected_cols, df_raw.copy()
            
            left_ctrl, right_panel = st.columns([1, 3.5])
            with left_ctrl:
                sel_filter_col = st.selectbox("🎯 快速筛选维度", ["不筛选"] + display_cols)
                num_cols = df_raw.select_dtypes(include=['number']).columns.tolist()
                selected_num_filters = st.multiselect("🔢 数值过滤字段", [c for c in num_cols if c in display_cols])
            
            with right_panel:
                if sel_filter_col != "不筛选":
                    u_vals = sorted(df_raw[sel_filter_col].dropna().unique().astype(str).tolist())
                    sel_vals = st.multiselect(f"选择 [{sel_filter_col}]", u_vals)
                    if sel_vals: df_final = df_final[df_final[sel_filter_col].astype(str).isin(sel_vals)]
                
                if selected_num_filters:
                    for col_name in selected_num_filters:
                        c_min, c_max = float(df_raw[col_name].min()), float(df_raw[col_name].max())
                        r_col1, r_col2, r_col3, r_col4, r_col5 = st.columns([1.2, 1, 1, 1, 1])
                        with r_col1: st.markdown(f"**{col_name}**")
                        v_min = r_col2.number_input("下限", value=c_min, key=f"min_{col_name}")
                        v_max = r_col3.number_input("上限", value=c_max, key=f"max_{col_name}")
                        q_input = r_col4.number_input("分位数%", 0, 100, 100, key=f"q_{col_name}")
                        q_actual = df_raw[col_name].quantile(q_input / 100.0)
                        with r_col5: st.info(f"≤ {q_actual:.2f}")
                        df_final = df_final[(df_final[col_name] >= v_min) & (df_final[col_name] <= min(v_max, q_actual))]

            s_c1, s_c2, _ = st.columns([1, 1, 2])
            s_c1.metric("筛选后行数", len(df_final))
            s_c2.metric("覆盖比例", f"{(len(df_final)/len(df_raw)*100):.1f}%")
            st.markdown('<div class="table-title">📋 数据样例前50条展示</div>', unsafe_allow_html=True)
            st.dataframe(df_final[display_cols].head(50), use_container_width=True)

        with tab2:
            current_cols = st.session_state.selected_cols
            if current_cols:
                # 1. 增加散点图选项
                chart_type = st.radio("✨ 图形形态", ["柱状图", "数据表格", "词云图", "折线图", "散点图", "饼图", "热力图"], horizontal=True)
                st.session_state.is_multi = st.toggle("开启多统计维度", value=st.session_state.get('is_multi', False))

                c1, c2, c3, c4, c5 = st.columns([1.5, 1.2, 1, 0.8, 1.5])
                g_by_cols = c1.multiselect("📌 统计维度(X)", current_cols, key="multi_x") if st.session_state.is_multi else [c1.selectbox("📌 统计维度(X)", current_cols, key="single_x")]
                y_target, meth = c2.selectbox("🎯 对象(Y)", current_cols, key="y_sel"), c3.selectbox("🧮 方法", ["计数", "求和", "平均值", "中位数"], key="meth_sel")
                t_n = c4.number_input("Top N (0为全部)", 0, value=15)

                custom_y_name = c5.text_input("🏷️ 数值名称", value=f"{y_target}-{meth}")
                color_theme = c5.selectbox("🎨 配色方案", ["默认蓝","商务蓝(纯色)", "柔和色", "深邃色", "鲜艳色", "灰度图"])

                if g_by_cols:
                    m_map = {"计数": "count", "求和": "sum", "平均值": "mean", "最小值": "min", "最大值": "max", "中位数": "median"}
                    y_label = f"{y_target}-{meth}"
                    t_n_val = t_n if t_n > 0 else 1000000 
                    
                    df_calc = df_final.copy()

                    # --- 数据处理分支：散点图 vs 其他聚合图表 ---
                    if chart_type == "散点图":
                        # 散点图直接取原始分布数据，不进行 groupby
                        y_label = y_target 
                        df_shown = df_calc[list(set(g_by_cols + [y_target]))].copy()
                        df_shown[y_target] = pd.to_numeric(df_shown[y_target], errors='coerce')
                        df_shown = df_shown.dropna(subset=[y_target]).head(t_n_val).reset_index(drop=True)
                        plot_x = "组合维度" if len(g_by_cols) > 1 else g_by_cols[0]
                        if len(g_by_cols) > 1: df_shown[plot_x] = df_shown[g_by_cols].astype(str).agg("-".join, axis=1)
                    else:
                        if meth == "计数":
                            df_agg = df_calc.groupby(g_by_cols).size().reset_index(name=y_label)
                        else:
                            df_calc[y_target] = pd.to_numeric(df_calc[y_target], errors='coerce')
                            df_agg = df_calc.groupby(g_by_cols)[y_target].agg(m_map[meth]).reset_index(name=y_label)
                        df_shown = df_agg.sort_values(y_label, ascending=False).head(t_n_val).reset_index(drop=True)
                        plot_x = "组合维度" if len(g_by_cols) > 1 else g_by_cols[0]
                        if len(g_by_cols) > 1: df_shown[plot_x] = df_shown[g_by_cols].astype(str).agg("-".join, axis=1)
                    
                    # --- 绘图配置 ---
                    current_fig = None
                    color_map = {"商务蓝(纯色)": ["#1f77b4"], "默认蓝": px.colors.sequential.Blues, "柔和色": px.colors.qualitative.Pastel, "深邃色": px.colors.qualitative.Prism, "鲜艳色": px.colors.qualitative.Vivid, "灰度图": px.colors.sequential.Greys}
                    color_seq = color_map.get(color_theme, px.colors.sequential.Blues)
                    labels_map = {y_label: custom_y_name if chart_type != "散点图" else y_target, plot_x: plot_x}

                    if chart_type == "数据表格": st.dataframe(df_shown.rename(columns={y_label: custom_y_name}), use_container_width=True)
                    elif chart_type == "词云图": 
                        current_fig = generate_professional_wordcloud(df_shown, g_by_cols[0], y_label)
                        st.pyplot(current_fig)
                    elif chart_type == "柱状图": 
                        current_fig = px.bar(df_shown, x=plot_x, y=y_label, color_discrete_sequence=color_seq if "纯色" in color_theme else None, color=y_label if "纯色" not in color_theme else None, color_continuous_scale=color_seq if "纯色" not in color_theme else None, text_auto=True, labels=labels_map)
                    elif chart_type == "折线图": 
                        current_fig = px.line(df_shown, x=plot_x, y=y_label, markers=True, labels=labels_map, color_discrete_sequence=color_seq if isinstance(color_seq, list) else [color_seq[-1]])
                    elif chart_type == "饼图": 
                        current_fig = px.pie(df_shown, names=plot_x, values=y_label, hole=0.3, labels=labels_map, color_discrete_sequence=color_seq if isinstance(color_seq, list) else color_seq[::-1])
                    elif chart_type == "散点图":
                        current_fig = px.scatter(df_shown, x=plot_x, y=y_label, color=y_label if "纯色" not in color_theme else None, color_discrete_sequence=color_seq if "纯色" in color_theme else None, color_continuous_scale=color_seq if "纯色" not in color_theme else None, labels=labels_map)
                        current_fig.update_traces(marker=dict(size=8))
                    elif chart_type == "热力图" and len(g_by_cols) >= 2: 
                        current_fig = px.density_heatmap(df_shown, x=g_by_cols[0], y=g_by_cols[1], z=y_label, color_continuous_scale=color_seq, text_auto=True, labels=labels_map)
                    elif chart_type == "热力图":
                        st.warning("⚠️ 热力图需要开启'多统计维度'并选择至少2个字段")

                    # --- 统一渲染逻辑 ---
                    if current_fig and hasattr(current_fig, "update_layout"):
                        current_fig.update_layout(template="plotly_white", title="", font_color="black", paper_bgcolor='white', coloraxis_showscale=False, legend_font_color="black")
                        current_fig.update_xaxes(title_font_color="black", tickfont_color="black", gridcolor="#f0f0f0", showline=True, linecolor='#e5e7eb')
                        current_fig.update_yaxes(title_font_color="black", tickfont_color="black", gridcolor="#f0f0f0", showline=True, linecolor='#e5e7eb')
                        if chart_type == "饼图": 
                            current_fig.update_traces(textfont_color="black", textinfo='label+percent', insidetextorientation='horizontal')
                            current_fig.update_layout(legend_x=1, legend_y=0.5, legend_xanchor="left", margin=dict(r=10, l=10))
                        st.plotly_chart(current_fig, use_container_width=True)

                    st.markdown("---")
                    with st.expander("📝 专家角色与提示词设定", expanded=False):
                        st.session_state.custom_prompt = st.text_area("提示词 (Prompt)", value=st.session_state.custom_prompt, height=150)
                    
                    rep_c1, rep_c2, rep_c3 = st.columns([1.2, 2.5, 1.2])
                    if rep_c1.button("🚀 生成专家研判", type="primary", use_container_width=True):
                        with st.spinner("专家研判生成中..."):
                            st.session_state.last_analysis = get_sh_expert_analysis(df_shown, g_by_cols, meth, client, model_id, st.session_state.custom_prompt, env_type)
                    
                    dim_str = "-".join(g_by_cols) if isinstance(g_by_cols, list) else str(g_by_cols)
                    default_frag_title = f"{dim_str}-{y_target}-{meth} 分析"
                    custom_frag_title = rep_c2.text_input("📝 片段标题", value=default_frag_title, label_visibility="collapsed", placeholder="给这个分析片段命名...")
                    
                    if st.session_state.last_analysis:
                        if rep_c3.button("➕ 暂存至报表", use_container_width=True):
                            st.session_state.report_list.append({
                                "title": custom_frag_title if custom_frag_title else default_frag_title,
                                "fig": current_fig, 
                                "text": st.session_state.last_analysis
                            })
                            st.toast("已加入清单")
                        st.markdown(f'<div class="report-box">{st.session_state.last_analysis}</div>', unsafe_allow_html=True)

                if st.session_state.report_list:
                    st.divider()
                    st.subheader(f"📑 汇总报表清单 ({len(st.session_state.report_list)})")
                    st.session_state.report_title = st.text_input("📝 自定义报告总标题", value=st.session_state.report_title)
                    
                    for i, item in enumerate(st.session_state.report_list):
                        with st.expander(f"片段 {i+1}: {item['title']}", expanded=False):
                            if item["fig"]:
                                try:
                                    if hasattr(item["fig"], "to_json"): st.plotly_chart(item["fig"], use_container_width=True, key=f"rep_chart_{i}")
                                    else: st.pyplot(item["fig"])
                                except: st.info("预览图表异常。")
                            st.write(item["text"])
                            if st.button(f"🗑️ 移除片段", key=f"del_{i}"):
                                st.session_state.report_list.pop(i); st.rerun()
                    
                    btn_col_word, btn_col_html = st.columns(2)
                    word_out = generate_word_report(st.session_state.report_title, st.session_state.report_list)
                    btn_col_word.download_button(
                        label="📥 下载 Word 研判报告 (含截图)",
                        data=word_out,
                        file_name=f"{st.session_state.report_title}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True
                    )

                    html_out = generate_html_report(st.session_state.report_title, st.session_state.report_list)
                    btn_col_html.download_button(
                        label="📥 下载 HTML 研判报告",
                        data=html_out,
                        file_name=f"{st.session_state.report_title}.html",
                        mime="text/html",
                        use_container_width=True
                    )

    except Exception as e: st.error(f"逻辑异常: {e}")